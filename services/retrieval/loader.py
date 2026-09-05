"""索引构建语料加载：六类源 → Node 列表（对话/任务/世界观/loading/情报/实体）。

对应 docs/v3-developer/04-检索与向量模型.md §3.4。文档解析与切分逻辑照搬
ai_engine/game_data_loader.py 的现有实现（业务行为不变），仅把 LlamaIndex
Document 换成 retrieval.store.Node；PDF 用 pypdf、DOCX 用 python-docx 直接解析。
"""

from __future__ import annotations

import json
import logging
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Callable, List

from services.game_data.paths import find_resources_directory
from services.retrieval.hybrid import cjk_tokenizer
from services.retrieval.store import Node, corpus_fingerprint

logger = logging.getLogger(__name__)

# 设定文档按章节/段落切分：识别标题行（用于按章节切分）
_LORE_HEADING_PATTERN = re.compile(
    r"^\s*(?:"
    r"#+\s+.+|"  # Markdown: # ## ###
    r"[一二三四五六七八九十百千]+[、．.]\s*.+|"  # 一、 二、
    r"\d+[.、]\s*.+|"  # 1. 2. 1、
    r"[（(][一二三四五六七八九十\d]+[)）]\s*.+"  # （一） (1)
    r")\s*$",
    re.MULTILINE,
)

# 情报 TXT 的 @@@X_Y@@@ 分节标记
_INTEL_SECTION_RE = re.compile(r"@@@\d+(?:_\d+)?@@@")

# 核心设定文档文件名标识（「重置知识库」的业务判定）
CORE_LORE_DOC_MARKER = "核心设定与世界合理性补足"

# 参与语料指纹的源目录（相对 resources 根）；任一文件 mtime/size 变化即触发重建
_FINGERPRINT_SOURCE_DIRS: tuple[str, ...] = (
    "data/dialogues",
    "data/task",
    "docs",
    "data/intelligence",
    "data/stages",
    "data/items",
)


def _resources_dir() -> Path:
    return find_resources_directory()


# ---------------------------------------------------------------------------
# 语料指纹
# ---------------------------------------------------------------------------

def compute_corpus_fingerprint(resources_dir: Path | None = None) -> str:
    """扫描全部语料源文件，产出 (相对路径, mtime_ns, size) 集合哈希。"""
    root = Path(resources_dir) if resources_dir else _resources_dir()
    states: list[tuple[str, int, int]] = []
    for rel_dir in _FINGERPRINT_SOURCE_DIRS:
        base = root / rel_dir
        if not base.exists():
            continue
        for f in base.rglob("*"):
            if f.is_file():
                try:
                    stat = f.stat()
                except OSError:
                    continue
                states.append((f.relative_to(root).as_posix(), stat.st_mtime_ns, stat.st_size))
    return corpus_fingerprint(states)


# ---------------------------------------------------------------------------
# 1. 日常对话 XML
# ---------------------------------------------------------------------------

def load_dialogue_nodes(resources_dir: Path | None = None) -> List[Node]:
    """读取 resources/data/dialogues 下的 NPC 日常对话 XML，每个 <Dialogue> 一条。"""
    root_dir = Path(resources_dir) if resources_dir else _resources_dir()
    dialogues_dir = root_dir / "data" / "dialogues"
    list_path = dialogues_dir / "list.xml"
    if not list_path.exists():
        logger.warning("未找到对话列表文件: %s", list_path)
        return []

    filenames = [
        (elem.text or "").strip()
        for elem in ET.parse(list_path).getroot().findall(".//items")
        if (elem.text or "").strip()
    ]

    nodes: List[Node] = []
    seq = 0
    for name in filenames:
        file_path = dialogues_dir / name
        if not file_path.exists():
            continue
        xml_root = ET.parse(file_path).getroot()
        file_level_name_elem = xml_root.find(".//Dialogues/Name")
        file_level_name = (
            file_level_name_elem.text.strip() if file_level_name_elem is not None else None
        )

        for dlg in xml_root.findall(".//Dialogues/Dialogue"):
            lines: List[str] = []
            character_key: str | None = None
            for sub in dlg.findall("./SubDialogue"):
                text_elem = sub.find("Text")
                text = (text_elem.text or "").strip() if text_elem is not None else ""
                if not text:
                    continue
                sub_name = (sub.find("Name").text or "").strip() if sub.find("Name") is not None else ""
                sub_char = (sub.find("Char").text or "").strip() if sub.find("Char") is not None else ""
                # 跳过玩家视角的台词（$PC / $PC_TITLE / $PC_CHAR 等）
                if sub_name == "$PC" or sub_char.startswith("$PC"):
                    continue
                # 角色标注使用 Name（角色名），缺省时回退文件级 Name 或 Char
                if character_key is None:
                    if sub_name:
                        character_key = sub_name
                    elif file_level_name:
                        character_key = file_level_name
                    elif sub_char:
                        character_key = sub_char.split("#")[0].strip()
                    else:
                        character_key = "Unknown"
                lines.append(text)

            if not lines or not character_key:
                continue
            seq += 1
            nodes.append(
                Node(
                    id=f"dialogue-{seq}",
                    text="\n".join(lines),
                    type="dialogue",
                    character=character_key.strip().lower(),  # 与检索端 npc_name 过滤一致
                    source_file=name,
                )
            )
    return nodes


# ---------------------------------------------------------------------------
# 2. 任务台词 JSON
# ---------------------------------------------------------------------------

def _is_player_dialogue_item(item: dict) -> bool:
    """判断对话条是否为玩家（$PC / $PC_TITLE / $PC_CHAR），此类不进入任务台词检索。"""
    if not isinstance(item, dict):
        return True
    name = str(item.get("name") or "").strip()
    title = str(item.get("title") or "").strip()
    char = str(item.get("char") or "").strip()
    if name == "$PC" or title == "$PC_TITLE":
        return True
    if char and (char == "$PC_CHAR" or char.startswith("$PC_CHAR#")):
        return True
    return False


def _task_character_from_item(item: dict) -> str | None:
    """从对话条取 NPC 角色名（Name 优先），规范化为小写；玩家条返回 None。"""
    if _is_player_dialogue_item(item):
        return None
    name = str(item.get("name") or "").strip()
    if not name:
        char = str(item.get("char") or "").strip()
        if char and not char.startswith("$PC"):
            name = char.split("#", maxsplit=1)[0].strip()
    return name.lower() if name else None


def load_task_nodes(resources_dir: Path | None = None) -> List[Node]:
    """任务台词：*tasks*.json + text/*.json，单条 NPC 台词一节点（guide 类打 task_source）。"""
    root_dir = Path(resources_dir) if resources_dir else _resources_dir()
    task_dir = root_dir / "data" / "task"
    text_dir = task_dir / "text"
    if not task_dir.exists() or not text_dir.exists():
        return []

    # 合并所有 text/*.json 的 key；preview_text.json 仅用于任务预览，显式跳过
    text_data: dict[str, object] = {}
    for jpath in sorted(text_dir.glob("*.json")):
        if jpath.stem == "preview_text":
            continue
        try:
            data = json.loads(jpath.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                text_data.update(data)
        except Exception as exc:
            logger.warning("跳过任务文本 %s: %s", jpath.name, exc)

    all_tasks: List[dict] = []
    for jpath in sorted(task_dir.glob("*tasks*.json")):
        try:
            data = json.loads(jpath.read_text(encoding="utf-8"))
            tasks = data.get("tasks") if isinstance(data, dict) else []
            if not isinstance(tasks, list):
                continue
            # 教学引导类仅在高分时才采用，避免与 NPC 形象弱关联时混入
            task_source = "guide" if "guide" in jpath.name.lower() else None
            for task in tasks:
                if not isinstance(task, dict):
                    continue
                t = dict(task)
                if task_source:
                    t["_task_source"] = task_source
                all_tasks.append(t)
        except Exception as exc:
            logger.warning("跳过任务配置 %s: %s", jpath.name, exc)

    nodes: List[Node] = []
    seq = 0
    for task in all_tasks:
        task_source: str | None = task.get("_task_source")
        for key in (task.get("get_conversation"), task.get("finish_conversation")):
            if not key:
                continue
            raw = text_data.get(str(key))
            if not isinstance(raw, list):
                continue
            for item in raw:
                if not isinstance(item, dict):
                    continue
                text = str(item.get("text") or "").strip()
                character = _task_character_from_item(item)
                if not text or not character:
                    continue
                seq += 1
                nodes.append(
                    Node(
                        id=f"task-{seq}",
                        text=text,
                        type="task",
                        character=character,
                        task_source=task_source,
                    )
                )
    return nodes


# ---------------------------------------------------------------------------
# 3. 情报 TXT（@@@X_Y@@@ 分节）
# ---------------------------------------------------------------------------

def load_intelligence_nodes(resources_dir: Path | None = None) -> List[Node]:
    root_dir = Path(resources_dir) if resources_dir else _resources_dir()
    intel_dir = root_dir / "data" / "intelligence"
    if not intel_dir.exists():
        logger.warning("情报目录不存在，跳过: %s", intel_dir)
        return []

    nodes: List[Node] = []
    for txt_file in sorted(intel_dir.glob("*.txt")):
        try:
            content = txt_file.read_text(encoding="utf-8").strip()
        except Exception:
            continue
        if len(content) < 10:
            continue
        source_name = txt_file.stem
        if _INTEL_SECTION_RE.search(content):
            for sec in _INTEL_SECTION_RE.split(content):
                sec = sec.strip()
                if len(sec) < 10:
                    continue
                nodes.append(
                    Node(id=f"intel-{len(nodes) + 1}", text=sec, type="intelligence", source_file=source_name)
                )
        else:
            nodes.append(
                Node(id=f"intel-{len(nodes) + 1}", text=content, type="intelligence", source_file=source_name)
            )
    return nodes


# ---------------------------------------------------------------------------
# 4. 世界观 PDF/DOCX（含按标题/段落/句子的 256/512 token 切分）
# ---------------------------------------------------------------------------

def _read_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def load_lore_nodes(resources_dir: Path | None = None) -> List[Node]:
    """读取 resources/docs 下的 PDF/DOCX 世界观文档，按文件名区分核心/补充设定。"""
    root_dir = Path(resources_dir) if resources_dir else _resources_dir()
    docs_dir = root_dir / "docs"
    if not docs_dir.exists():
        logger.warning("世界观设定目录不存在，跳过: %s", docs_dir)
        return []

    matching_files = [
        f for f in docs_dir.iterdir()
        if f.is_file() and f.suffix.lower() in (".pdf", ".docx")
    ]
    if not matching_files:
        logger.warning("docs 目录中无 PDF/DOCX 文件，跳过世界观文档加载")
        return []

    nodes: List[Node] = []
    for f in sorted(matching_files):
        try:
            text = _read_pdf_text(f) if f.suffix.lower() == ".pdf" else _read_docx_text(f)
        except Exception as exc:
            logger.warning("解析世界观文档 %s 失败，跳过: %s", f.name, exc)
            continue
        if not (text or "").strip():
            continue
        doc_type = "world_lore" if CORE_LORE_DOC_MARKER in f.stem else "supplementary_lore"
        nodes.append(
            Node(id=f"lore-raw-{len(nodes) + 1}", text=text, type=doc_type, source_file=f.name)
        )
    return nodes


def _split_by_headings(text: str) -> List[str]:
    """按标题行将文本拆成多个章节（Markdown #、一、二、、1. 2.、（一）等）。"""
    if not (text or "").strip():
        return []
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    if not blocks:
        return [text.strip()] if text.strip() else []

    sections: List[str] = []
    current: List[str] = []
    for block in blocks:
        first_line = block.split("\n")[0] if "\n" in block else block
        is_heading = bool(_LORE_HEADING_PATTERN.match(first_line.strip()))
        if is_heading and current:
            sections.append("\n\n".join(current))
            current = [block]
        else:
            current.append(block)
    if current:
        sections.append("\n\n".join(current))
    return sections


def _split_sentences(text: str) -> List[str]:
    """按。！？；分句，保留边界完整（不 mid-sentence 截断）。"""
    if not text or not text.strip():
        return []
    parts = re.split(r"([。！？；])", text)
    sentences: List[str] = []
    buf = ""
    for p in parts:
        buf += p
        if p.strip() in "。！？；" and buf.strip():
            sentences.append(buf.strip())
            buf = ""
    if buf.strip():
        sentences.append(buf.strip())
    return sentences


_LINE_END_PUNCTUATION = set("。！？；，、．·.?!;:：,，！？；")


def _ends_with_punctuation(s: str) -> bool:
    t = (s or "").rstrip()
    return bool(t) and t[-1] in _LINE_END_PUNCTUATION


def _normalize_pdf_soft_line_breaks(text: str) -> str:
    """只把「不以标点结尾的换行」拼接到下一行，保留真正的段落大换行（PDF 行宽换行修正）。"""
    if not text or not text.strip():
        return text
    lines = text.split("\n")
    result: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            result.append("")
            i += 1
            continue
        buf = line
        j = i + 1
        while j < len(lines) and lines[j].strip() and not _ends_with_punctuation(buf):
            buf = buf + lines[j]
            j += 1
        result.append(buf)
        i = j
    return "\n".join(result)


def chunk_lore_nodes(lore_nodes: List[Node], tokenizer: Callable[[str], List[str]]) -> List[Node]:
    """设定文档按章节/段落切分并应用 256/512 token 规则，避免断句（沿用旧实现）。"""
    CHUNK_TARGET = 256
    CHUNK_MAX_SINGLE = 512

    def token_count(t: str) -> int:
        return len(tokenizer(t)) if t and t.strip() else 0

    result: List[Node] = []
    for raw in lore_nodes:
        text = (raw.text or "").strip()
        if not text:
            continue
        # 仅对 PDF 做软换行整合；Word 等已有正确段落结构保持原样
        if (raw.source_file or "").lower().endswith(".pdf"):
            text = _normalize_pdf_soft_line_breaks(text)

        sections = _split_by_headings(text) or [text]
        candidates: List[str] = []

        for sec in sections:
            sec = sec.strip()
            if not sec:
                continue
            if token_count(sec) <= CHUNK_MAX_SINGLE:
                candidates.append(sec)
                continue
            # 章节 > 512：按段落拆
            for para in (p.strip() for p in sec.split("\n\n")):
                if not para:
                    continue
                if token_count(para) <= CHUNK_MAX_SINGLE:
                    candidates.append(para)
                    continue
                # 段落 > 512：按句号分句后成块
                buf = ""
                for s in _split_sentences(para):
                    if token_count(s) > CHUNK_MAX_SINGLE:
                        if buf.strip():
                            candidates.append(buf.strip())
                            buf = ""
                        candidates.append(s)
                        continue
                    merged = (buf + "\n" + s).strip() if buf else s
                    if token_count(merged) <= CHUNK_TARGET:
                        buf = merged
                        continue
                    if buf.strip():
                        candidates.append(buf.strip())
                    buf = s
                if buf.strip():
                    candidates.append(buf.strip())

        # 合并过短的候选块到约 CHUNK_TARGET
        i = 0
        while i < len(candidates):
            chunk = candidates[i]
            if token_count(chunk) >= CHUNK_TARGET:
                result.append(Node(id="", text=chunk, type=raw.type, source_file=raw.source_file))
                i += 1
                continue
            merged = chunk
            j = i + 1
            while j < len(candidates):
                merged_next = (merged + "\n\n" + candidates[j]).strip()
                if token_count(merged_next) > CHUNK_MAX_SINGLE:
                    break
                merged = merged_next
                j += 1
                if token_count(merged) >= CHUNK_TARGET:
                    break
            result.append(Node(id="", text=merged, type=raw.type, source_file=raw.source_file))
            i = j

    for k, node in enumerate(result):
        node.id = f"lore-{k + 1}"
    return result


# ---------------------------------------------------------------------------
# 5. loading 文案 XML
# ---------------------------------------------------------------------------

def load_loading_nodes(resources_dir: Path | None = None) -> List[Node]:
    root_dir = Path(resources_dir) if resources_dir else _resources_dir()
    xml_path = root_dir / "data" / "stages" / "loading_data.xml"
    if not xml_path.exists():
        logger.warning("未找到 loading 文本文件，跳过: %s", xml_path)
        return []
    try:
        root = ET.parse(xml_path).getroot()
    except Exception as exc:
        logger.warning("解析 loading_data.xml 出错，跳过: %s", exc)
        return []

    nodes: List[Node] = []
    for group in root.findall(".//LoadingText/Group"):
        region_elem = group.find("Region")
        unlock_elem = group.find("Unlock")
        region = (region_elem.text or "").strip() if region_elem is not None else ""
        unlock_raw = (unlock_elem.text or "").strip() if unlock_elem is not None else ""
        for text_elem in group.findall("Text"):
            text = (text_elem.text or "").strip() if text_elem is not None else ""
            if not text:
                continue
            nodes.append(
                Node(
                    id=f"loading-{len(nodes) + 1}",
                    text=text,
                    type="loading_lore",
                    source_file="loading_data.xml",
                    region=region or None,
                    unlock=unlock_raw or None,
                )
            )
    return nodes


# ---------------------------------------------------------------------------
# 6. 物品/关卡结构化实体（整实体一条）
# ---------------------------------------------------------------------------

def load_game_entity_nodes(resources_dir: Path | None = None) -> List[Node]:
    """物品与关卡各建一条向量节点；构建索引时直接实例化 Registry 读盘。"""
    from services.game_data.item_registry import ItemRegistry
    from services.game_data.paths import get_game_data_root
    from services.game_data.stage_registry import StageRegistry
    from services.game_entity_prompts import format_item_embedding_text, format_stage_embedding_text

    root = get_game_data_root()
    items = ItemRegistry(data_root=root)
    items.load()
    stages = StageRegistry(data_root=root)
    stages.load()

    nodes: List[Node] = []
    for it in items.items:
        text = format_item_embedding_text(it)
        if not (text or "").strip():
            continue
        nodes.append(
            Node(id=f"game-item-{len(nodes) + 1}", text=text, type="game_item", item_name=it.name)
        )
    for si in stages._stage_infos.values():
        text = format_stage_embedding_text(si)
        nodes.append(
            Node(
                id=f"game-stage-{len(nodes) + 1}",
                text=text,
                type="game_stage",
                stage_area=si.area,
                stage_name=si.name,
                entity_key=f"{si.area}::{si.name}",
            )
        )
    return nodes


# ---------------------------------------------------------------------------
# 汇总
# ---------------------------------------------------------------------------

def load_corpus(resources_dir: Path | None = None) -> List[Node]:
    """加载全部六类语料（世界观文档单独按 256/512 token 规则切块）。"""
    root_dir = Path(resources_dir) if resources_dir else _resources_dir()

    nodes: List[Node] = []
    nodes.extend(load_dialogue_nodes(root_dir))
    nodes.extend(load_task_nodes(root_dir))
    try:
        lore_raw = load_lore_nodes(root_dir)
    except Exception as exc:
        logger.warning("加载世界观文档时出错，跳过: %s", exc)
        lore_raw = []
    try:
        nodes.extend(load_loading_nodes(root_dir))
    except Exception as exc:
        logger.warning("加载 loading 文本时出错，跳过: %s", exc)
    try:
        nodes.extend(load_intelligence_nodes(root_dir))
    except Exception as exc:
        logger.warning("加载情报文件时出错，跳过: %s", exc)
    try:
        nodes.extend(load_game_entity_nodes(root_dir))
    except Exception as exc:
        logger.warning("加载游戏实体向量文档时出错，跳过: %s", exc)

    if lore_raw:
        lore_chunks = chunk_lore_nodes(lore_raw, cjk_tokenizer)
        logger.info("设定文档切分: %d 个原文档 -> %d 个块", len(lore_raw), len(lore_chunks))
        nodes.extend(lore_chunks)

    logger.info(
        "语料加载完成: %d 条节点（对话=%d 任务=%d 世界观块=%d loading=%d 情报=%d 实体=%d）",
        len(nodes),
        sum(1 for n in nodes if n.type == "dialogue"),
        sum(1 for n in nodes if n.type == "task"),
        sum(1 for n in nodes if n.type in ("world_lore", "supplementary_lore")),
        sum(1 for n in nodes if n.type == "loading_lore"),
        sum(1 for n in nodes if n.type == "intelligence"),
        sum(1 for n in nodes if n.type in ("game_item", "game_stage")),
    )
    if not nodes:
        raise ValueError("没有加载到任何语料节点，无法构建检索索引。")
    return nodes


def has_core_lore_document(resources_dir: Path | None = None) -> bool:
    """resources/docs 下是否存在「核心设定与世界合理性补足」文档（重置知识库的业务判定）。"""
    try:
        root_dir = Path(resources_dir) if resources_dir else _resources_dir()
    except FileNotFoundError:
        return False
    docs_dir = root_dir / "docs"
    if not docs_dir.exists():
        return False
    for f in docs_dir.iterdir():
        if f.is_file() and f.suffix.lower() in (".pdf", ".docx") and CORE_LORE_DOC_MARKER in f.stem:
            return True
    return False
